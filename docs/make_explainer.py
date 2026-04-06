"""
make_explainer.py — generates CSV_Analyzer_Explainer.docx from scratch.
Run:   python docs/make_explainer.py
Output: docs/CSV_Analyzer_Explainer.docx

Sections:
  1.  What is this project?
  2.  What problem does it solve?
  3.  How does it work? (pipeline)
  4.  How is the code organised?
  5.  How does the program store data?
  6.  What does the output look like?          ← includes student highlights
  7.  What did the Coursera courses teach?
  8.  Glossary
  9.  The formulas explained (mean, median, min, max, count)
  10. Why StudentID appears & how names were added
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ------------------------------------------------------------------ #
# Helpers
# ------------------------------------------------------------------ #

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def heading(doc, text, level=1, color="1f4e79"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = 'Arial'
    return p

def body(doc, text, italic=False, bold=False, color="333333", size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.italic = italic
    run.font.bold   = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def table(doc, rows_data, col_widths,
          header_bg="1f4e79", alt_bg="dae3f3", alt2_bg="FFFFFF"):
    """Build a styled table. First row is always the header (dark bg)."""
    n_cols = len(col_widths)
    tbl = doc.add_table(rows=len(rows_data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        bg = header_bg if is_header else (alt_bg if r_idx % 2 == 1 else alt2_bg)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(col_widths[c_idx] / 1440)
            set_cell_bg(cell, bg)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name  = 'Arial'
            run.font.size  = Pt(10)
            run.font.bold  = is_header
            run.font.color.rgb = (RGBColor(0xFF, 0xFF, 0xFF)
                                  if is_header else RGBColor(0x33, 0x33, 0x33))
    return tbl

def sp(doc):
    """Add a blank spacing paragraph."""
    doc.add_paragraph()

# ------------------------------------------------------------------ #
# Build the document
# ------------------------------------------------------------------ #

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ================================================================== #
# COVER
# ================================================================== #
sp(doc)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("CSV Data Analyzer — Plain-English Explainer")
tr.font.name = 'Arial'; tr.font.size = Pt(20)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A guide for interviewers, professors, family & friends")
sr.font.name = 'Arial'; sr.font.size = Pt(12)
sr.font.italic = True
sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
sp(doc)

# ================================================================== #
# SECTION 1 — What is this project?
# ================================================================== #
heading(doc, "1.  What is this project?")
body(doc,
    "Imagine you have a big Excel-like file — a CSV file — filled with rows of "
    "numbers: student marks, sales figures, temperatures, anything. "
    "This project is a small computer program written in the C programming language "
    "that reads that file and automatically tells you the key statistics for every "
    "numeric column: the average, the middle value, the lowest number, and the "
    "highest number.")
body(doc,
    "You run it from the command line (the black terminal window), give it the "
    "file name, and it prints a neat table of results in seconds. You can also "
    "save those results to a text file.")

# ================================================================== #
# SECTION 2 — What problem does it solve?
# ================================================================== #
heading(doc, "2.  What problem does it solve?")
body(doc,
    "Doing statistics by hand is tedious and error-prone. While tools like Excel "
    "exist, this project shows that the same task can be done with a small, "
    "self-written program — one that runs on any computer, requires no internet, "
    "and processes even very large files (100,000+ rows) quickly.")
body(doc,
    "The real educational value: building this project from scratch teaches the "
    "programmer how a computer stores data in memory, how it reads files line by "
    "line, and how it performs calculations — all things that happen invisibly "
    "inside every app you use every day.")

# ================================================================== #
# SECTION 3 — How does it work?
# ================================================================== #
heading(doc, "3.  How does it work? (Step-by-step)")
body(doc, "The program follows a four-step pipeline every time it runs:",
     italic=True, color="555555")
sp(doc)

table(doc, [
    ["Step",  "What happens",                         "Real-world analogy"],
    ["1",
     "Open & Read\nThe program opens the CSV file and reads it line by line.",
     "Like flipping through pages of a notebook."],
    ["2",
     "Parse\nSplits each line at the commas, converts text like \"88.5\" "
     "into a real number the computer can calculate with.",
     "Like cutting a sentence into individual words."],
    ["3",
     "Compute Stats\nFor each numeric column calculates:\n"
     "  Mean, Median, Min, Max\n"
     "Also records which student (row label) produced the min and max.",
     "Like a calculator working through your data."],
    ["4",
     "Display & Export\nResults are printed as a neat table on screen "
     "and optionally saved to a .txt file.",
     "Like printing a report from Excel."],
], col_widths=[600, 4200, 3000])
sp(doc)

# ================================================================== #
# SECTION 4 — Code organisation
# ================================================================== #
heading(doc, "4.  How is the code organised?")
body(doc,
    "Instead of writing everything in one giant file, the code is split into "
    "five separate files, each with one job. Think of it like a kitchen — "
    "there is a prep station, a cooking station, and a plating station.")
sp(doc)

table(doc, [
    ["File",          "Job (one sentence)",                           "Analogy"],
    ["main.c",
     "Starts the program, reads the filename you typed, calls the "
     "other modules in order.",
     "The manager."],
    ["csv_parser.c",
     "Opens the file, reads every row, converts text into numbers, "
     "saves the Name column as labels.",
     "The data-entry clerk."],
    ["stats.c",
     "Loops through the numbers and calculates mean, median, min, max. "
     "Also records which student scored the min and max.",
     "The mathematician."],
    ["display.c",
     "Takes the results and prints them as a formatted table; "
     "writes to file if asked.",
     "The report writer."],
    ["utils.c",
     "Small helper tools: safe memory allocation, text trimming, "
     "number validation.",
     "The toolbox."],
], col_widths=[1080, 4200, 2520],
   header_bg="375623", alt_bg="e2efda")
sp(doc)

# ================================================================== #
# SECTION 5 — Memory layout
# ================================================================== #
heading(doc, "5.  How does the program store data?")
body(doc,
    "After reading the CSV file, all numbers live in the computer's RAM "
    "(temporary working memory) as a grid — rows going down, columns going "
    "across. The student names are stored separately as a list of strings "
    "(one per row). Once the program finishes, everything is freed.")
sp(doc)
body(doc, "Simplified memory layout after reading a 3-row, 3-column CSV:",
     italic=True, color="555555")

table(doc, [
    ["",       "Col 0: Maths", "Col 1: Physics", "Col 2: Biology"],
    ["Row 0",  "88.5",         "76.0",           "83.5"],
    ["Row 1",  "72.0",         "68.5",           "79.0"],
    ["Row 2",  "95.0",         "89.0",           "91.0"],
], col_widths=[900, 2100, 2100, 2100],
   header_bg="833c00", alt_bg="fbe5d6")

body(doc, "Name labels stored separately: Row 0 = Alice, Row 1 = Bob, Row 2 = Charlie",
     italic=True, color="555555")
sp(doc)

# ================================================================== #
# SECTION 6 — What the output looks like  (UPDATED with highlights)
# ================================================================== #
heading(doc, "6.  What does the output look like?")
body(doc,
    "After running the program on the sample student-marks CSV (15 students), "
    "the terminal shows two tables.")
sp(doc)

body(doc, "Table 1 — Statistical summary (one row per numeric column):", bold=True)

table(doc, [
    ["Column",      "Mean",    "Median",   "Min",    "Max"],
    ["StudentID",   "108.00",  "108.00",   "101.00", "115.00"],
    ["Maths",       "79.03",   "83.00",    "55.00",  "95.00"],
    ["Physics",     "75.33",   "76.00",    "55.00",  "94.50"],
    ["Chemistry",   "78.63",   "80.50",    "57.50",  "93.50"],
    ["Biology",     "78.67",   "80.50",    "58.50",  "94.00"],
    ["Attendance",  "86.80",   "88.00",    "70.00",  "98.00"],
], col_widths=[1620, 1530, 1530, 1530, 1530])
sp(doc)

body(doc, "Table 2 — Student highlights (who scored the lowest and highest):", bold=True)

table(doc, [
    ["Column",      "Min score",  "Lowest scored by",   "Max score", "Highest scored by"],
    ["Maths",       "55.00",      "Hank",               "95.00",     "Charlie"],
    ["Physics",     "55.00",      "Diana",              "94.50",     "Grace"],
    ["Chemistry",   "57.50",      "Hank",               "93.50",     "Karen"],
    ["Biology",     "58.50",      "Diana",              "94.00",     "Olivia"],
    ["Attendance",  "70.00",      "Hank",               "98.00",     "Charlie"],
], col_widths=[1440, 1200, 2160, 1200, 2160])
sp(doc)

body(doc,
    "Note: StudentID also appears in the stats table because 101, 102 ... are "
    "numbers. The program cannot know they are labels rather than measurements. "
    "See Section 10 for a full explanation.",
     italic=True, color="833c00")
sp(doc)

# ================================================================== #
# SECTION 7 — Coursera courses
# ================================================================== #
heading(doc, "7.  What did the Coursera courses teach?")
body(doc,
    "Two courses from the University of California, Santa Cruz "
    "(completed July 2024) provided the foundation for this project:")
sp(doc)

table(doc, [
    ["Course",                                   "Key concept learned",
     "Used in this project"],
    ["C for Everyone, Part 1:\nProgramming Fundamentals",
     "Variables, loops, functions,\npointers, memory (malloc/free), file I/O",
     "Reading the CSV, storing numbers, malloc/free"],
    ["C for Everyone, Part 2:\nStructured Programming",
     "Structs, multi-file programs,\nsorting algorithms, modular design",
     "CSVData & ColumnStats structs,\nqsort for median, 5-file structure"],
], col_widths=[2160, 2880, 2880])
sp(doc)

# ================================================================== #
# SECTION 8 — Glossary
# ================================================================== #
heading(doc, "8.  Glossary — words you might hear")

table(doc, [
    ["Term",            "Plain-English meaning"],
    ["CSV file",
     "A simple spreadsheet saved as plain text with commas separating values — "
     "like a list you could type in Notepad."],
    ["C language",
     "A programming language from the 1970s that is very fast and used inside "
     "operating systems, cars, and satellites."],
    ["Compiler (GCC)",
     "A tool that translates human-readable code into instructions the computer "
     "processor understands."],
    ["RAM / Memory",
     "The computer's short-term working space — like a desk. Data lives here "
     "while the program runs, then disappears."],
    ["malloc / free",
     "Commands that ask the computer for a chunk of memory, then give it back "
     "when done — like borrowing and returning library books."],
    ["Struct",
     "A container that groups related pieces of data — like a form with fields "
     "(Name, Age, Score)."],
    ["qsort",
     "A built-in C sorting function used to arrange values in order so the "
     "middle one (median) can be found."],
    ["NaN",
     "Not a Number — a special marker the program stores when a cell contains "
     "text (like a name). Stats functions skip NaN cells."],
    ["CLI / Terminal",
     "The black text window where you type commands directly — used to run "
     "this program."],
], col_widths=[1500, 6300],
   header_bg="404040", alt_bg="f2f2f2")
sp(doc)

# ================================================================== #
# SECTION 9 — Formulas explained with actual computed values
# ================================================================== #
heading(doc, "9.  The Formulas — What exactly is being calculated?")
body(doc,
    "For every numeric column the program computes four statistics. "
    "Here is what each formula means, how it is computed, and the actual "
    "result from our sample data (15 students, Maths column).")
sp(doc)

# 9a — Full formula table
table(doc, [
    ["Statistic",      "Formula",
     "How the program does it",
     "Actual result\n(Maths, 15 students)"],

    ["Mean\n(Average)",
     "Sum of all values\n÷ number of values",
     "Loop through every value, keep a running total.\n"
     "Divide total by count at the end.",
     "Sum = 1185.50\nCount = 15\n\nMean = 1185.50 ÷ 15\n= 79.03"],

    ["Median\n(Middle value)",
     "Sort all values.\nPick the middle one.\n\n"
     "If count is even:\n  average the two middle values.\n"
     "If count is odd:\n  take the single middle value.",
     "1. Copy the values into a temp array.\n"
     "2. Sort using qsort().\n"
     "3. Pick index [count/2] if odd,\n"
     "   or average [n/2-1] and [n/2] if even.",
     "Sorted: 55, 61.5, 62, 69.5, 72,\n74, 78.5, [83], 84, 87.5,\n"
     "88.5, 90.5, 91, 93.5, 95\n\n8th value (odd, 15 items)\n= 83.00"],

    ["Min\n(Lowest value)",
     "Scan all values.\nReturn the smallest.",
     "Start with first value as 'current min'.\n"
     "Compare each next value — if smaller,\nupdate current min.",
     "Lowest mark = 55.00\n(scored by Hank)"],

    ["Max\n(Highest value)",
     "Scan all values.\nReturn the largest.",
     "Start with first value as 'current max'.\n"
     "Compare each next value — if larger,\nupdate current max.",
     "Highest mark = 95.00\n(scored by Charlie)"],

    ["Count\n(Valid rows)",
     "Count how many cells\ncontained a real number\n(skips text & blank cells).",
     "During parsing, text cells are stored as NaN.\n"
     "Count = number of non-NaN values in the column.",
     "All 15 students have\na Maths mark.\nCount = 15"],
], col_widths=[1050, 1650, 2700, 2400])
sp(doc)

# 9b — Summary of all computed values for all columns
body(doc,
    "Computed values for all columns in sample.csv (15 students):",
     bold=True)
sp(doc)

table(doc, [
    ["Column",      "Mean",    "Median",  "Min",    "Max",   "Count"],
    ["StudentID",   "108.00",  "108.00",  "101.00", "115.00","15"],
    ["Maths",       "79.03",   "83.00",   "55.00",  "95.00", "15"],
    ["Physics",     "75.33",   "76.00",   "55.00",  "94.50", "15"],
    ["Chemistry",   "78.63",   "80.50",   "57.50",  "93.50", "15"],
    ["Biology",     "78.67",   "80.50",   "58.50",  "94.00", "15"],
    ["Attendance",  "86.80",   "88.00",   "70.00",  "98.00", "15"],
], col_widths=[1440, 1170, 1170, 1170, 1170, 1080])

body(doc,
    "Tip: Mean and Median being close together (e.g. Maths: 79.03 vs 83.00) "
    "means the data is fairly evenly spread. A big gap between them would "
    "suggest a few very high or very low outliers pulling the average.",
     italic=True, color="555555")
sp(doc)

# ================================================================== #
# SECTION 10 — StudentID & student names
# ================================================================== #
heading(doc, "10.  \"Why does StudentID appear?\" — and how names were added")

heading(doc, "Why StudentID gets statistics", level=2, color="375623")
body(doc,
    "The program does not know the meaning of a column — it only checks whether "
    "the values look like numbers. StudentID values (101, 102 ...) pass that "
    "check, so the program computes stats for them.")
body(doc,
    "The results are mathematically correct but meaningless in real life. "
    "An ID is a label, not a measurement. This is a well-known problem in "
    "data analysis called 'treating identifiers as variables'.")
sp(doc)

table(doc, [
    ["What the program sees",   "What it means in real life"],
    ["StudentID is a column of numbers",
     "It is just a roll number (a label). It has no meaningful average."],
    ["Mean StudentID = 108.00",
     "Meaningless — the average of roll numbers tells you nothing useful."],
    ["Min = 101 (Alice), Max = 115 (Olivia)",
     "Just means Alice was entered first and Olivia last — not about performance."],
], col_widths=[3600, 4800],
   header_bg="833c00", alt_bg="fbe5d6")
sp(doc)

heading(doc, "How student names appear in the output", level=2, color="1f4e79")
body(doc,
    "The program automatically detects the first column that contains text "
    "values (the 'Name' column). It saves those strings in a separate list "
    "in memory — one name per row. After computing statistics, it traces back "
    "which row (which student) produced the minimum and maximum value in each "
    "subject, and looks up their name.")
sp(doc)

table(doc, [
    ["Step", "What happens in the code"],
    ["1. Detect label column",
     "During CSV parsing, the first non-numeric column found is saved as the "
     "'label column' (csv->label_col_idx). For our file, that is the Name column."],
    ["2. Save the strings",
     "Each name (Alice, Bob ...) is stored in a char** array called row_labels[], "
     "one slot per row — separate from the numeric data grid."],
    ["3. Find min/max row",
     "After computing the minimum value for a column, the program scans the "
     "data to find which row number contains that exact value."],
    ["4. Look up the name",
     "It uses that row number to look up row_labels[row_number], giving the "
     "student's name, which is then stored in ColumnStats as min_by / max_by."],
    ["5. Display",
     "display.c checks if min_by is not NULL and, if so, prints the "
     "Student Highlights table below the main stats table."],
], col_widths=[1620, 6180],
   header_bg="1f4e79", alt_bg="dae3f3")
sp(doc)

# ================================================================== #
# Footer
# ================================================================== #
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run(
    "Project: CSV Data Analyzer in C  |  Language: C (C11)  |  "
    "Courses: C for Everyone Part 1 & 2, UC Santa Cruz — Coursera (July 2024)")
nr.font.name   = 'Arial'
nr.font.size   = Pt(9)
nr.font.italic = True
nr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ================================================================== #
# Save
# ================================================================== #
out_path = r"D:\Coursera C Project\docs\CSV_Analyzer_Explainer.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print("Sections included:")
for i, s in enumerate([
    "What is this project?",
    "What problem does it solve?",
    "How does it work?",
    "How is the code organised?",
    "How does the program store data?",
    "What does the output look like? (with student highlights)",
    "What did the Coursera courses teach?",
    "Glossary",
    "The formulas explained (mean, median, min, max, count) with actual values",
    "Why StudentID appears & how names were added",
], 1):
    print(f"  {i:2d}. {s}")
